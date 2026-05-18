"""
ENGINE: Maintenance Memorandum DOCX Exporter

Експортує MaintenanceResult у DOCX-меморандум.
Підтримка UA/EN через параметр language.
"""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from engine.maintenance_calculator import MaintenanceResult


# ═══════════════════════════════════════════════════════════════════
# КОЛЬОРИ
# ═══════════════════════════════════════════════════════════════════
COLOR_HEADER = RGBColor(0x2E, 0x75, 0xB6)
COLOR_H2 = RGBColor(0x1F, 0x4E, 0x79)
COLOR_ACCENT = RGBColor(0xC0, 0x39, 0x2B)
COLOR_HIGHLIGHT = RGBColor(0xFF, 0xA5, 0x00)
BG_LIGHT = "F0F4F8"
BG_HEADER = "D5E8F0"
BG_HIGHLIGHT = "FFF8E7"
BG_DISCOUNT = "FDEDEC"
TEXT_GRAY = RGBColor(0x55, 0x55, 0x55)
DARK_TEXT = RGBColor(0x33, 0x33, 0x33)


# ═══════════════════════════════════════════════════════════════════
# ЛОКАЛІЗАЦІЯ
# ═══════════════════════════════════════════════════════════════════
L10N = {
    "ua": {
        "title": "МЕМОРАНДУМ\nрозрахунку вартості технічного обслуговування СПЗ",
        "title_with_mfr": "МЕМОРАНДУМ\nрозрахунку вартості ТО СПЗ\n({mfr_name})",
        "section_1": "1. Дані про об'єкт",
        "section_2": "2. Прийняті розрахункові показники",
        "section_3": "3. Фінансові параметри",
        "section_4": "4. Розрахунок витрат часу",
        "section_5": "5. Розрахунок собівартості",
        "section_6": "6. Розрахунок ціни послуги",
        "section_7": "7. Підсумок",
        "section_8": "8. Примітки та умови",
        # Section 1
        "param": "Параметр",
        "value": "Значення",
        "area": "Загальна площа захищуваної території",
        "composition": "Склад СПЗ",
        "distance": "Відстань до об'єкта від офісу",
        "subcontract": "Пультове спостереження",
        "subcontract_yes": "виконує підрядна організація",
        "manufacturer": "Виробник системи",
        # Section 2
        "indicator": "Показник",
        "basis": "Підстава",
        "base_time": "Базовий час планового ТО",
        "complexity_k": "Коефіцієнт складності системи",
        "n_false": "Прогнозована к-ть хибних спрацювань",
        "n_damages": "Прогнозована к-ть пошкоджень",
        "n_visits": "Кількість планових візитів",
        "auto_from_area": "auto для площі",
        "from_composition": "за складом системи",
        "from_mfr_level": "за рівнем виробника",
        "typical": "типове значення",
        "optimal": "оптимально",
        "month_per": "/міс",
        # Section 3
        "salary": "Середня з/п кваліфікованого виконавця",
        "tax": "Сумарний податок (нарахування на оплату праці)",
        "work_hours": "Робочих годин на місяць",
        "rate": "Брутто-ставка години роботи (з податками)",
        "fuel_cons": "Витрата пального",
        "fuel_price": "Вартість пального",
        "amortization": "Амортизація автомобіля",
        "amortization_eq_fuel": "= вартості пального",
        "transport_km": "Загальна вартість 1 км пробігу",
        "admin_rate": "Адміністративні витрати (від ОП)",
        "markup": "Планова націнка (markup) до собівартості",
        "subcontract_cost": "Підрядна послуга (пультове)",
        # Section 4
        "time_formula": "Формула: T = T_планове + T_дорога + T_хибні + T_пошкодження",
        "time_item": "Стаття витрат часу",
        "time_formula_col": "Формула / Обчислення",
        "hours": "Год",
        "t_planned": "Планове ТО (з урахуванням складності)",
        "t_travel": "Дорога на планові візити",
        "t_false_alarms": "Реагування на хибні (з дорогою)",
        "t_damages": "Усунення пошкоджень (з дорогою)",
        "t_total": "РАЗОМ людино-годин на місяць",
        "sum": "сумарно",
        # Section 5
        "cost_item": "Стаття витрат",
        "cost_uah": "Сума, грн",
        "cost_labor": "Оплата праці кваліфікованого виконавця (з податками)",
        "cost_transport": "Транспортні витрати",
        "cost_parts": "Запчастини / витратні матеріали",
        "cost_admin": "Адміністративні витрати",
        "cost_own_total": "СОБІВАРТІСТЬ ВЛАСНИХ РОБІТ",
        # Section 6
        "price_item": "Складова ціни",
        "price_own": "Власні роботи (з націнкою {markup_pct}%)",
        "subcontract_pass": "Підрядне пультове спостереження (прохід)",
        "by_subcontract": "за рахунком підрядника",
        "price_calculated": "РОЗРАХУНКОВА ЦІНА",
        "discount_label": "Стратегічна знижка",
        "discount_formula": "−{pct}% від розрахункової",
        "price_final": "ЦІНА, ЗАПРОПОНОВАНА КЛІЄНТУ",
        "minus": "розрахункова − знижка",
        "own_plus_sub": "власні + підряд",
        # Section 7
        "month_total": "Вартість ТО на місяць",
        "year_total": "Вартість ТО на рік",
        "uah": "грн",
        # Section 8 notes
        "notes": [
            ("Структура націнки",
             "Націнка {markup_pct}% (markup) застосована лише до власних витрат. "
             "Підрядна послуга пультового спостереження передається замовнику за "
             "фактичною вартістю без націнки. Це забезпечує прозорість і "
             "конкурентоспроможність ціни."),
            ("Орієнтовний характер показників",
             "Кількість хибних спрацювань та пошкоджень — це прогнозні значення на "
             "основі типової статистики для об'єктів аналогічного розміру та складу. "
             "За підсумками 3–6 місяців експлуатації може бути проведено уточнення "
             "розрахунку зі скоригованими показниками."),
            ("Що не входить у вартість",
             "У розрахунок НЕ включено: вартість заміни обладнання, що вийшло з ладу "
             "не з вини виконавця; модернізацію системи; екстрені виїзди у вихідні та "
             "святкові дні; ПДВ (нараховується додатково за чинним законодавством)."),
            ("Періодичність ТО",
             "Згідно з ДСТУ-Н CEN/TS 54-14:2021 технічне обслуговування виконується "
             "щомісячно з повним регламентним обслуговуванням щоквартально. У "
             "розрахунку враховано планові візити на місяць."),
        ],
        # Generated
        "generated": "Згенеровано: FireCompare v{ver} | {date}",
    },
    "en": {
        "title": "MEMORANDUM\nfire alarm system maintenance cost calculation",
        "title_with_mfr": "MEMORANDUM\nFAS maintenance cost calculation\n({mfr_name})",
        "section_1": "1. Object data",
        "section_2": "2. Accepted calculation parameters",
        "section_3": "3. Financial parameters",
        "section_4": "4. Time costs calculation",
        "section_5": "5. Cost breakdown",
        "section_6": "6. Price calculation",
        "section_7": "7. Summary",
        "section_8": "8. Notes and conditions",
        # Section 1
        "param": "Parameter",
        "value": "Value",
        "area": "Total protected area",
        "composition": "FAS composition",
        "distance": "Distance to object from office",
        "subcontract": "Monitoring service",
        "subcontract_yes": "performed by subcontractor",
        "manufacturer": "System manufacturer",
        # Section 2
        "indicator": "Indicator",
        "basis": "Basis",
        "base_time": "Base planned maintenance time",
        "complexity_k": "System complexity coefficient",
        "n_false": "Projected false alarms",
        "n_damages": "Projected damages",
        "n_visits": "Planned visits",
        "auto_from_area": "auto for area",
        "from_composition": "by system composition",
        "from_mfr_level": "by manufacturer level",
        "typical": "typical value",
        "optimal": "optimal",
        "month_per": "/month",
        # Section 3
        "salary": "Qualified executor avg. salary",
        "tax": "Total tax (labor charges)",
        "work_hours": "Work hours per month",
        "rate": "Gross hourly rate (with taxes)",
        "fuel_cons": "Fuel consumption",
        "fuel_price": "Fuel price",
        "amortization": "Vehicle amortization",
        "amortization_eq_fuel": "= fuel cost",
        "transport_km": "Total cost per 1 km",
        "admin_rate": "Administrative costs (of labor)",
        "markup": "Planned markup to own cost",
        "subcontract_cost": "Subcontractor service (monitoring)",
        # Section 4
        "time_formula": "Formula: T = T_planned + T_travel + T_false_alarms + T_damages",
        "time_item": "Time cost item",
        "time_formula_col": "Formula / Calculation",
        "hours": "Hours",
        "t_planned": "Planned maintenance (with complexity factor)",
        "t_travel": "Travel for planned visits",
        "t_false_alarms": "Response to false alarms (with travel)",
        "t_damages": "Damage repair (with travel)",
        "t_total": "TOTAL man-hours per month",
        "sum": "sum",
        # Section 5
        "cost_item": "Cost item",
        "cost_uah": "Amount, UAH",
        "cost_labor": "Qualified executor labor (with taxes)",
        "cost_transport": "Transport costs",
        "cost_parts": "Spare parts / consumables",
        "cost_admin": "Administrative costs",
        "cost_own_total": "OWN WORK COST",
        # Section 6
        "price_item": "Price component",
        "price_own": "Own work (with {markup_pct}% markup)",
        "subcontract_pass": "Subcontracted monitoring (pass-through)",
        "by_subcontract": "by subcontractor invoice",
        "price_calculated": "CALCULATED PRICE",
        "discount_label": "Strategic discount",
        "discount_formula": "−{pct}% of calculated",
        "price_final": "PRICE OFFERED TO CLIENT",
        "minus": "calculated − discount",
        "own_plus_sub": "own + subcontract",
        # Section 7
        "month_total": "Monthly maintenance cost",
        "year_total": "Annual maintenance cost",
        "uah": "UAH",
        # Section 8 notes
        "notes": [
            ("Markup structure",
             "The {markup_pct}% markup is applied only to own costs. The subcontracted "
             "monitoring service is passed through to the client at the actual cost "
             "without markup. This ensures transparency and competitive pricing."),
            ("Indicative nature of figures",
             "The number of false alarms and damages are projected values based on "
             "typical statistics for objects of similar size and composition. After "
             "3-6 months of operation, the calculation can be refined with actual data."),
            ("What is not included",
             "NOT included: replacement of equipment failed not by executor's fault; "
             "system modernization; emergency visits on weekends and holidays; VAT "
             "(charged additionally per current legislation)."),
            ("Maintenance frequency",
             "Per DSTU-N CEN/TS 54-14:2021, maintenance is performed monthly with full "
             "scheduled servicing quarterly. The calculation includes planned visits "
             "per month."),
        ],
        # Generated
        "generated": "Generated: FireCompare v{ver} | {date}",
    },
}


def L(lang: str, key: str, **kwargs) -> str:
    """Translate key with optional formatting"""
    template = L10N.get(lang, L10N["ua"]).get(key, key)
    if kwargs:
        try:
            return template.format(**kwargs)
        except (KeyError, IndexError):
            return template
    return template


# ═══════════════════════════════════════════════════════════════════
# DOCX HELPERS
# ═══════════════════════════════════════════════════════════════════


def _set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def _add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(18)
            run.font.color.rgb = COLOR_HEADER
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = COLOR_H2
    return h


def _add_para(doc, text, *, bold=False, italic=False, size=11, color=DARK_TEXT, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return p


def _add_table(doc, headers, rows, *, highlight_rows=None, col_widths=None):
    if highlight_rows is None:
        highlight_rows = {}
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell, BG_HEADER)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = COLOR_H2
    
    for r_i, row in enumerate(rows):
        highlight = highlight_rows.get(r_i)
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if highlight:
                _set_cell_shading(cell, highlight)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            if highlight:
                run.bold = True
    
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    
    return table


# ═══════════════════════════════════════════════════════════════════
# ГОЛОВНА ФУНКЦІЯ
# ═══════════════════════════════════════════════════════════════════


def export_maintenance_memo_to_docx(
    result: MaintenanceResult,
    output_path: str,
    language: str = "ua",
) -> str:
    """
    Експортує меморандум розрахунку ТО у DOCX.
    
    Args:
        result: результат calculate_maintenance()
        output_path: куди зберегти
        language: "ua" або "en"
    
    Returns:
        шлях до створеного файлу
    """
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    
    bd = result.breakdown
    params = result.params
    composition = params.composition
    markup_pct = int(params.markup * 100)
    tax_pct = int(params.tax_rate * 100)
    admin_pct = int(params.admin_rate * 100)
    
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
    
    # ═══ ТИТУЛ ═══
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if result.manufacturer_name:
        title_text = L(language, "title_with_mfr", mfr_name=result.manufacturer_name)
    else:
        title_text = L(language, "title")
    run = title_para.add_run(title_text)
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = COLOR_HEADER
    
    # Subtitle
    _add_para(
        doc,
        L(language, "generated", ver="0.3.0", date=datetime.now().strftime("%Y-%m-%d %H:%M")),
        italic=True, color=TEXT_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, size=10,
    )
    doc.add_paragraph()
    
    # ═══ SECTION 1: Object data ═══
    _add_heading(doc, L(language, "section_1"), level=2)
    
    s1_rows = [
        [L(language, "area"), f"{params.object_area_m2:,.0f} m²"],
        [L(language, "composition"), composition.composition_label(lang=language)],
        [L(language, "distance"), f"{params.distance_km} km"],
    ]
    if composition.has_monitoring_subcontract:
        s1_rows.append([L(language, "subcontract"), L(language, "subcontract_yes")])
    if result.manufacturer_name:
        s1_rows.append([L(language, "manufacturer"), result.manufacturer_name])
    
    _add_table(
        doc,
        headers=[L(language, "param"), L(language, "value")],
        rows=s1_rows,
        col_widths=[7, 9],
    )
    doc.add_paragraph()
    
    # ═══ SECTION 2: Calculation parameters ═══
    _add_heading(doc, L(language, "section_2"), level=2)
    
    base_hours_basis = (
        f"{L(language, 'auto_from_area')} {params.object_area_m2:,.0f} m²"
        if params.base_planned_hours is None
        else "manual"
    )
    
    s2_rows = [
        [L(language, "base_time"), f"{bd.t_planned / bd.complexity_k:.1f} h{L(language, 'month_per')}", base_hours_basis],
        [L(language, "complexity_k"), f"{bd.complexity_k:.2f}", L(language, "from_composition")],
        [L(language, "n_false"), 
         f"{params.n_false_alarms_month or 'auto'}{L(language, 'month_per')}",
         L(language, "from_mfr_level") if params.n_false_alarms_month is None else L(language, "typical")],
        [L(language, "n_damages"), f"{params.n_damages_month}{L(language, 'month_per')}",
         L(language, "typical")],
        [L(language, "n_visits"), f"{params.n_planned_visits}{L(language, 'month_per')}",
         L(language, "optimal")],
    ]
    _add_table(
        doc,
        headers=[L(language, "indicator"), L(language, "value"), L(language, "basis")],
        rows=s2_rows,
        col_widths=[6, 4, 6],
    )
    doc.add_paragraph()
    
    # ═══ SECTION 3: Financial parameters ═══
    _add_heading(doc, L(language, "section_3"), level=2)
    
    fuel_cost_per_km = params.fuel_consumption_l_100km / 100 * params.fuel_price_uah_l
    
    s3_rows = [
        [L(language, "salary"), f"{params.salary_uah:,.0f} {L(language, 'uah')}{L(language, 'month_per')}"],
        [L(language, "tax"), f"{tax_pct}%"],
        [L(language, "work_hours"), f"{params.work_hours_month}"],
        [L(language, "rate"), f"{bd.rate_per_hour:.2f} {L(language, 'uah')}/h"],
        [L(language, "fuel_cons"), f"{params.fuel_consumption_l_100km} l/100 km"],
        [L(language, "fuel_price"), f"{params.fuel_price_uah_l} {L(language, 'uah')}/l"],
        [L(language, "amortization"), 
         f"{L(language, 'amortization_eq_fuel')} ({fuel_cost_per_km:.0f} {L(language, 'uah')}/km)"],
        [L(language, "transport_km"), f"{bd.transport_per_km} {L(language, 'uah')}/km"],
        [L(language, "admin_rate"), f"{admin_pct}%"],
        [L(language, "markup"), f"{markup_pct}%"],
    ]
    if composition.has_monitoring_subcontract:
        s3_rows.append([
            L(language, "subcontract_cost"),
            f"{composition.subcontract_monitoring_uah:,.0f} {L(language, 'uah')}{L(language, 'month_per')}",
        ])
    
    _add_table(
        doc,
        headers=[L(language, "param"), L(language, "value")],
        rows=s3_rows,
        col_widths=[10, 6],
    )
    doc.add_page_break()
    
    # ═══ SECTION 4: Time calculation ═══
    _add_heading(doc, L(language, "section_4"), level=2)
    _add_para(doc, L(language, "time_formula"))
    
    n_false_actual = params.n_false_alarms_month if params.n_false_alarms_month is not None else (
        params.n_false_alarms_month  # буде з default
    )
    
    one_trip_h = 2 * params.distance_km / 50
    s4_rows = [
        [L(language, "t_planned"), 
         f"× {bd.complexity_k:.2f}",
         f"{bd.t_planned:.2f}"],
        [L(language, "t_travel"),
         f"{params.n_planned_visits} × 2 × {params.distance_km} / 50",
         f"{bd.t_travel_planned:.2f}"],
        [L(language, "t_false_alarms"),
         f"× (1.5 + {one_trip_h:.2f})",
         f"{bd.t_false_alarms:.2f}"],
        [L(language, "t_damages"),
         f"{params.n_damages_month} × (3.0 + {one_trip_h:.2f})",
         f"{bd.t_damages:.2f}"],
        [L(language, "t_total"), L(language, "sum"), f"{bd.t_total:.2f}"],
    ]
    _add_table(
        doc,
        headers=[L(language, "time_item"), L(language, "time_formula_col"), L(language, "hours")],
        rows=s4_rows,
        col_widths=[7, 6, 3],
        highlight_rows={4: BG_HIGHLIGHT},
    )
    doc.add_paragraph()
    
    # ═══ SECTION 5: Cost ═══
    _add_heading(doc, L(language, "section_5"), level=2)
    
    s5_rows = [
        [L(language, "cost_labor"),
         f"{bd.t_total:.2f} h × {bd.rate_per_hour:.2f}",
         f"{bd.cost_labor:,.2f}"],
        [L(language, "cost_transport"),
         f"{bd.total_km:.1f} km × {bd.transport_per_km}",
         f"{bd.cost_transport:,.2f}"],
        [L(language, "cost_parts"),
         f"{params.n_damages_month} × parts",
         f"{bd.cost_parts:,.2f}"],
        [L(language, "cost_admin"),
         f"labor × {admin_pct}%",
         f"{bd.cost_admin:,.2f}"],
        [L(language, "cost_own_total"), L(language, "sum"), f"{bd.cost_own_total:,.2f}"],
    ]
    _add_table(
        doc,
        headers=[L(language, "cost_item"), L(language, "time_formula_col"), L(language, "cost_uah")],
        rows=s5_rows,
        col_widths=[7, 6, 3],
        highlight_rows={4: BG_HIGHLIGHT},
    )
    doc.add_paragraph()
    
    # ═══ SECTION 6: Price ═══
    _add_heading(doc, L(language, "section_6"), level=2)
    
    s6_rows = [
        [L(language, "price_own", markup_pct=markup_pct),
         f"{bd.cost_own_total:,.2f} × {1 + params.markup}",
         f"{bd.price_own_calculated:,.2f}"],
    ]
    if bd.subcontract_pass_through > 0:
        s6_rows.append([
            L(language, "subcontract_pass"),
            L(language, "by_subcontract"),
            f"{bd.subcontract_pass_through:,.2f}",
        ])
    s6_rows.append([
        L(language, "price_calculated"),
        L(language, "own_plus_sub"),
        f"{bd.price_calculated_total:,.2f}",
    ])
    
    highlight = {len(s6_rows) - 1: BG_LIGHT}
    
    if bd.discount_uah > 0:
        s6_rows.append([
            L(language, "discount_label"),
            L(language, "discount_formula", pct=f"{params.strategic_discount_pct:.1f}"),
            f"−{bd.discount_uah:,.2f}",
        ])
        highlight[len(s6_rows) - 1] = BG_DISCOUNT
    
    s6_rows.append([
        L(language, "price_final"),
        L(language, "minus") if bd.discount_uah > 0 else L(language, "own_plus_sub"),
        f"{bd.price_final_month:,.2f}",
    ])
    highlight[len(s6_rows) - 1] = BG_HIGHLIGHT
    
    _add_table(
        doc,
        headers=[L(language, "price_item"), L(language, "time_formula_col"), L(language, "cost_uah")],
        rows=s6_rows,
        col_widths=[7, 6, 3],
        highlight_rows=highlight,
    )
    doc.add_paragraph()
    
    # ═══ SECTION 7: Summary ═══
    _add_heading(doc, L(language, "section_7"), level=2)
    
    summary = doc.add_table(rows=2, cols=2)
    summary.style = 'Light Grid Accent 1'
    
    for cell, text in zip(
        [summary.cell(0, 0), summary.cell(0, 1),
         summary.cell(1, 0), summary.cell(1, 1)],
        [L(language, "month_total"),
         f"{bd.price_final_month:,.2f} {L(language, 'uah')}",
         L(language, "year_total"),
         f"{bd.price_final_year:,.2f} {L(language, 'uah')}"],
    ):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell, BG_HIGHLIGHT)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(13)
        run.bold = True
        run.font.color.rgb = COLOR_H2
    
    doc.add_paragraph()
    
    # ═══ SECTION 8: Notes ═══
    _add_heading(doc, L(language, "section_8"), level=2)
    
    for title_text, body in L10N[language]["notes"]:
        # Підставити markup_pct у текст
        body_formatted = body.format(markup_pct=markup_pct)
        p = doc.add_paragraph()
        run = p.add_run(f"{title_text}. ")
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.bold = True
        run.font.color.rgb = COLOR_H2
        run = p.add_run(body_formatted)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = DARK_TEXT
    
    doc.save(str(output))
    return str(output)
