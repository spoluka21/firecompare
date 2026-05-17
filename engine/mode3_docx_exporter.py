"""
MODE 3 — DOCX EXPORTER (через python-docx)

Експортує Mode3Memo у формат DOCX без зовнішніх залежностей.
Замість Node.js docx-js використовуємо python-docx — це робить
розгортання на Streamlit Cloud простим (тільки pip-залежності).
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from engine.mode3_memo import Mode3Memo, MemoSection


# ═══════════════════════════════════════════════════════════════════
# КОЛЬОРОВА СХЕМА (відповідає Cofem-стилю UI)
# ═══════════════════════════════════════════════════════════════════
COLOR_HEADER_BLUE = RGBColor(0x2E, 0x75, 0xB6)
COLOR_SECTION_BLUE = RGBColor(0x1F, 0x4E, 0x79)
COLOR_TABLE_HEADER_BG = "D5E8F0"  # світло-блакитний для заголовків таблиць
COLOR_TEXT_BODY = RGBColor(0x33, 0x33, 0x33)


# ═══════════════════════════════════════════════════════════════════
# ДОПОМІЖНІ ФУНКЦІЇ
# ═══════════════════════════════════════════════════════════════════


def _set_cell_shading(cell, fill_hex: str):
    """Заливка комірки таблиці кольором"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)


def _add_heading(doc: Document, text: str, level: int = 1):
    """Додає заголовок з потрібним стилем"""
    heading = doc.add_heading(text, level=level)
    
    # Налаштовуємо колір і шрифт
    for run in heading.runs:
        run.font.name = "Arial"
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = COLOR_HEADER_BLUE
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = COLOR_SECTION_BLUE
    
    return heading


def _add_paragraph(doc: Document, text: str, italic: bool = False, alignment=None):
    """Додає звичайний параграф"""
    para = doc.add_paragraph()
    if alignment:
        para.alignment = alignment
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEXT_BODY
    if italic:
        run.italic = True
    return para


def _add_bullet(doc: Document, text: str):
    """Додає буллет-список"""
    para = doc.add_paragraph(text, style='List Bullet')
    for run in para.runs:
        run.font.name = "Arial"
        run.font.size = Pt(11)
    return para


def _add_table(doc: Document, table_data: list[list[str]]):
    """Додає таблицю з даних. Перший рядок — заголовок."""
    if not table_data or len(table_data) == 0:
        return None
    
    rows_count = len(table_data)
    cols_count = len(table_data[0])
    
    table = doc.add_table(rows=rows_count, cols=cols_count)
    table.style = 'Light Grid Accent 1'
    
    for row_idx, row_data in enumerate(table_data):
        is_header = row_idx == 0
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # Заливка для заголовка
            if is_header:
                _set_cell_shading(cell, COLOR_TABLE_HEADER_BG)
            
            # Текст
            cell_para = cell.paragraphs[0]
            cell_para.clear()
            run = cell_para.add_run(str(cell_text))
            run.font.name = "Arial"
            run.font.size = Pt(10)
            if is_header:
                run.bold = True
                run.font.color.rgb = COLOR_SECTION_BLUE
    
    return table


def _render_section(doc: Document, section: MemoSection):
    """Малює один розділ меморандуму"""
    _add_heading(doc, section.title, level=2)
    
    # Текст параграфів
    for para_text in section.content:
        _add_paragraph(doc, para_text)
    
    # Таблиця
    if section.table_data:
        _add_table(doc, section.table_data)
        # відступ після таблиці
        doc.add_paragraph()
    
    # Буллети
    for bullet in section.bullet_points:
        if bullet == "---":
            doc.add_paragraph()  # роздільник
            continue
        _add_bullet(doc, bullet)


# ═══════════════════════════════════════════════════════════════════
# ГОЛОВНА ФУНКЦІЯ
# ═══════════════════════════════════════════════════════════════════


def export_memo_to_docx(memo: Mode3Memo, output_path: str) -> str:
    """
    Експортує меморандум у DOCX через python-docx.
    
    Повертає шлях до створеного файлу.
    """
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    
    doc = Document()
    
    # Налаштовуємо поля сторінки (1 inch = 2.54 cm)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # ── Титул ──
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(memo.title)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(18)
    title_run.bold = True
    title_run.font.color.rgb = COLOR_HEADER_BLUE
    
    # Підзаголовок
    _add_paragraph(doc, memo.subtitle, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()  # відступ після титулу
    
    # ── Розділи ──
    for section in [
        memo.section_1_profile,
        memo.section_2_compliance,
        memo.section_3_technical,
        memo.section_4_economics,
        memo.section_5_position,
        memo.section_6_strengths_weaknesses,
        memo.section_7_summary,
    ]:
        _render_section(doc, section)
        doc.add_paragraph()  # відступ між розділами
    
    # ── Зберегти ──
    doc.save(str(output))
    return str(output)
